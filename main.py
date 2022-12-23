import requests
from bs4 import BeautifulSoup
import time
import concurrent.futures
import threading
import os
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink




# document.add_heading('Test', 0)

input_string = input('Search?.... ')
search_query = 'https://hbr.org/search?N=0&Ntt=' + '+'.join(input_string.split(' '))
# search_query = 'https://hbr.org/search?N=0&Ntt=multiple+business+location'

start = time.time()

document = Document()

document.add_heading(input_string, 0)

page = requests.get(search_query)
soup = BeautifulSoup(page.content, 'html.parser')
# print(soup.prettify())


# finding how many pages HBR has + applying less than 100 page rule. for now!
for index, item in enumerate(soup.find_all('h3')[1].strings):
    if index == 0:
        number = [x for x in item if x!=',' and x!=' ']
        number = int("".join(number))

number = int(number/10)
print(str(number) + ' Pages has found.')

if number > 100:
    number = 100
    print('Results of 100 pages will be saved.')


search_query  = search_query + '&loaded=' + str(number)
page = requests.get(search_query)
soup = BeautifulSoup(page.content, 'html.parser')


temp = soup.find_all('search-stream')[0].find_all('stream-list')[1].find_all('stream-item')
# print(*temp, sep='\n\n\n')

counter = 0
for item in temp:
    if item['data-list-price'] == '0':
        link = 'https://hbr.org/' + item.find('a')['href']
        title = item['data-title']
        summary = item['data-summary']
        counter += 1

        h = document.add_heading('', level=1)
        add_hyperlink(h, title, link)
        p = document.add_paragraph(summary)


print('\n\n' + str(counter) + ' Free articles has been saved.' + '\n')


document.save('/Users/Reza/Desktop/' + input_string + '.docx')

end = time.time()

# print('time: ' + str((end - start)/60))


